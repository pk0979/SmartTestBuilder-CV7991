import streamlit as st
import pandas as pd
import random
import io
import json
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- CẤU HÌNH TRANG ---
st.set_page_config(page_title="EduTest Pro - CV7991 (Form 2025)", page_icon="🎓", layout="wide")

st.markdown("""
<style>
    .stMetric { background-color: #f0f2f6; padding: 10px; border-radius: 10px; }
    .stTabs [data-baseweb="tab-list"] { gap: 24px; }
    .stTabs [data-baseweb="tab"] { height: 50px; white-space: pre-wrap; background-color: #ffffff; border-radius: 4px 4px 0 0; padding-top: 10px; padding-bottom: 10px; }
    .stTabs [aria-selected="true"] { background-color: #e6f0ff; border-bottom: 3px solid #0066cc; font-weight: bold;}
</style>
""", unsafe_allow_html=True)

if 'db_df' not in st.session_state:
    st.session_state.db_df = pd.DataFrame(columns=['Chuong', 'Muc_do', 'Loai_cau_hoi', 'Noi_dung', 'A', 'B', 'C', 'D', 'Dap_an_dung'])

# --- HÀM TIỆN ÍCH CHO WORD ---
def style_text(paragraph, text, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return run

def set_cell_margins(cell, top=50, bottom=50, start=50, end=50):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in [("top", top), ("bottom", bottom), ("left", start), ("right", end)]:
        node = OxmlElement(f'w:{m[0]}')
        node.set(qn('w:w'), str(m[1]))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# --- HÀM GỌI GEMINI AI (XỬ LÝ CÔNG THỨC TOÁN/HÓA) ---
def generate_questions_with_ai(api_key, subject, chapter, q_tn, q_ds, q_tl):
    prompt = f"""
    Bạn là chuyên gia giáo dục tại Việt Nam. Hãy soạn câu hỏi môn {subject}, phần "{chapter}" theo Form 2025.
    Số lượng: {q_tn} câu Trắc nghiệm, {q_ds} câu Đúng/Sai (Mỗi câu 4 ý a,b,c,d), {q_tl} câu Tự luận.
    
    ĐẶC BIỆT QUAN TRỌNG VỚI TOÁN/HÓA HỌC:
    - TUYỆT ĐỐI KHÔNG dùng mã LaTeX (không dùng $, $$, \\frac, _, ^).
    - BẮT BUỘC dùng ký tự Unicode để viết chỉ số trên/dưới. 
    - Ví dụ Hóa học ghi: H₂O, CO₂, SO₄²⁻, Cu²⁺. Toán học ghi: x², y³, aₙ. Nhiệt độ ghi: t°.
    
    Quy tắc JSON trả về (TUYỆT ĐỐI KHÔNG chứa markdown ```json):
    [
        {{"Chuong": "{chapter}", "Muc_do": "NB", "Loai_cau_hoi": "Trắc nghiệm", "Noi_dung": "...", "A": "...", "B": "...", "C": "...", "D": "...", "Dap_an_dung": "A"}},
        {{"Chuong": "{chapter}", "Muc_do": "VD", "Loai_cau_hoi": "Đúng/Sai", "Noi_dung": "...", "A": "Ý a...", "B": "Ý b...", "C": "Ý c...", "D": "Ý d...", "Dap_an_dung": "Đ, S, Đ, S"}},
        {{"Chuong": "{chapter}", "Muc_do": "VDC", "Loai_cau_hoi": "Tự luận", "Noi_dung": "...", "A": "", "B": "", "C": "", "D": "", "Dap_an_dung": "Đáp án..."}}
    ]
    """
    
    headers = {'Content-Type': 'application/json'}
    data = {"contents": [{"parts":[{"text": prompt}]}], "generationConfig": {"response_mime_type": "application/json"}}
    models = ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-1.5-flash", "gemini-1.5-flash-8b", "gemini-1.5-pro"]
    
    last_error = ""
    for model in models:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
        try:
            response = requests.post(url, headers=headers, json=data)
            if response.status_code == 200:
                result = response.json()['candidates'][0]['content']['parts'][0]['text']
                result = result.replace('```json', '').replace('```', '').strip()
                return json.loads(result)
            elif response.status_code == 404:
                continue
            else:
                last_error = response.text
        except Exception as e:
            last_error = str(e)
            continue
    st.error(f"❌ Lỗi kết nối AI: {last_error}")
    return None

def shuffle_question(row):
    loai = str(row.get('Loai_cau_hoi', 'Trắc nghiệm')).strip()
    if loai != 'Trắc nghiệm':
        return {
            'Noi_dung': row['Noi_dung'], 'Options': {'A': row.get('A',''), 'B': row.get('B',''), 'C': row.get('C',''), 'D': row.get('D','')},
            'Correct': row['Dap_an_dung'], 'Muc_do': row['Muc_do'], 'Loai': loai, 'Chuong': row.get('Chuong', 'Chung')
        }
    options = [('A', str(row['A'])), ('B', str(row['B'])), ('C', str(row['C'])), ('D', str(row['D']))]
    correct_content = str(row[row['Dap_an_dung']])
    random.shuffle(options)
    new_options, new_correct_label, labels = {}, "", ['A', 'B', 'C', 'D']
    for i in range(4):
        new_options[labels[i]] = options[i][1]
        if options[i][1] == correct_content: new_correct_label = labels[i]
    return {
        'Noi_dung': row['Noi_dung'], 'Options': new_options, 'Correct': new_correct_label,
        'Muc_do': row['Muc_do'], 'Loai': loai, 'Chuong': row.get('Chuong', 'Chung')
    }

# --- HÀM XUẤT WORD CHUẨN FORM CV7991 MỚI ---
def export_full_exam(exam_list, info):
    doc = Document()
    
    # Định dạng font mặc định toàn văn bản
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # === TRANG 1: MA TRẬN VÀ ĐẶC TẢ ===
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_text(p_head, f"{info['coquan'].upper()}\n", bold=True)
    style_text(p_head, f"{info['school'].upper()}\n\n", bold=True)
    style_text(p_head, f"MA TRẬN VÀ ĐẶC TẢ ĐỀ KIỂM TRA\n", bold=True)
    style_text(p_head, f"MÔN {info['subject'].upper()}\nNĂM HỌC 2025 - 2026\n", bold=True)

    # I. Mục đích
    doc.add_paragraph("I. MỤC ĐÍCH KIỂM TRA").runs[0].bold = True
    doc.add_paragraph("- Kiểm tra kiến thức đã học theo phân phối chương trình.")
    doc.add_paragraph("- HS biết vận dụng kiến thức đã học để làm bài kiểm tra.")
    
    # II. Hình thức
    doc.add_paragraph("II. HÌNH THỨC KIỂM TRA").runs[0].bold = True
    doc.add_paragraph(f"- Thời gian làm bài: {info['time']} phút.")
    doc.add_paragraph("- Hình thức kiểm tra: Kết hợp giữa trắc nghiệm và tự luận.")
    
    # Thống kê câu hỏi để vẽ ma trận
    stats = {}
    for q in exam_list:
        ch = q['Chuong']
        loai = q['Loai']
        md = q['Muc_do']
        # Gom nhóm mức độ theo Biết(B), Hiểu(H), Vận dụng (VD+VDC)
        if md == 'NB': md_key = 'B'
        elif md == 'TH': md_key = 'H'
        else: md_key = 'VD'
        
        if ch not in stats:
            stats[ch] = {'TNKQ':{'B':0, 'H':0, 'VD':0}, 'DS':{'B':0, 'H':0, 'VD':0}, 'TL':{'B':0, 'H':0, 'VD':0}}
        
        if loai == 'Trắc nghiệm': stats[ch]['TNKQ'][md_key] += 1
        elif loai == 'Đúng/Sai': stats[ch]['DS'][md_key] += 1
        else: stats[ch]['TL'][md_key] += 1

    # III. Ma trận
    doc.add_paragraph("III. MA TRẬN ĐẶC TẢ ĐỀ KIỂM TRA").runs[0].bold = True
    doc.add_paragraph("1. Ma trận").runs[0].bold = True
    
    # Bảng Ma Trận 14 cột phức tạp
    mt = doc.add_table(rows=3 + len(stats) + 2, cols=14)
    mt.style = 'Table Grid'
    
    # Gộp ô Tiêu đề
    mt.cell(0,0).merge(mt.cell(2,0)); mt.cell(0,0).text = "TT"
    mt.cell(0,1).merge(mt.cell(2,1)); mt.cell(0,1).text = "Chương/Chủ đề"
    mt.cell(0,2).merge(mt.cell(0,10)); mt.cell(0,2).text = "Mức độ đánh giá"
    mt.cell(0,11).merge(mt.cell(2,12)); mt.cell(0,11).text = "Tổng"
    mt.cell(0,13).merge(mt.cell(2,13)); mt.cell(0,13).text = "Tỉ lệ %"
    
    # Gộp ô Dòng 2
    mt.cell(1,2).merge(mt.cell(1,4)); mt.cell(1,2).text = "TNKQ"
    mt.cell(1,5).merge(mt.cell(1,7)); mt.cell(1,5).text = "Đúng/Sai"
    mt.cell(1,8).merge(mt.cell(1,10)); mt.cell(1,8).text = "Tự luận"
    
    # Điền B, H, VD
    labels_bhv = ["B", "H", "VD"] * 3
    for i in range(9):
        mt.cell(2, 2+i).text = labels_bhv[i]
        
    # Căn giữa Header
    for r in range(3):
        for c in range(14):
            p = mt.cell(r,c).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs: p.runs[0].bold = True

    # Đổ dữ liệu
    r_idx = 3
    t_tn_b = t_tn_h = t_tn_vd = 0
    t_ds_b = t_ds_h = t_ds_vd = 0
    t_tl_b = t_tl_h = t_tl_vd = 0
    
    for idx, (ch, data) in enumerate(stats.items()):
        mt.cell(r_idx, 0).text = str(idx + 1)
        mt.cell(r_idx, 1).text = str(ch)
        
        # TNKQ
        mt.cell(r_idx, 2).text = str(data['TNKQ']['B']) if data['TNKQ']['B']>0 else ""
        mt.cell(r_idx, 3).text = str(data['TNKQ']['H']) if data['TNKQ']['H']>0 else ""
        mt.cell(r_idx, 4).text = str(data['TNKQ']['VD']) if data['TNKQ']['VD']>0 else ""
        t_tn_b += data['TNKQ']['B']; t_tn_h += data['TNKQ']['H']; t_tn_vd += data['TNKQ']['VD']
        
        # DS
        mt.cell(r_idx, 5).text = f"{data['DS']['B']} ý" if data['DS']['B']>0 else ""
        mt.cell(r_idx, 6).text = f"{data['DS']['H']} ý" if data['DS']['H']>0 else ""
        mt.cell(r_idx, 7).text = f"{data['DS']['VD']} ý" if data['DS']['VD']>0 else ""
        t_ds_b += data['DS']['B']; t_ds_h += data['DS']['H']; t_ds_vd += data['DS']['VD']
        
        # TL
        mt.cell(r_idx, 8).text = str(data['TL']['B']) if data['TL']['B']>0 else ""
        mt.cell(r_idx, 9).text = str(data['TL']['H']) if data['TL']['H']>0 else ""
        mt.cell(r_idx, 10).text = str(data['TL']['VD']) if data['TL']['VD']>0 else ""
        t_tl_b += data['TL']['B']; t_tl_h += data['TL']['H']; t_tl_vd += data['TL']['VD']
        
        r_idx += 1
        
    # Dòng Tổng
    mt.cell(r_idx, 0).merge(mt.cell(r_idx, 1)); mt.cell(r_idx, 0).text = "Tổng số câu"
    mt.cell(r_idx, 2).text = str(t_tn_b); mt.cell(r_idx, 3).text = str(t_tn_h); mt.cell(r_idx, 4).text = str(t_tn_vd)
    mt.cell(r_idx, 5).text = str(t_ds_b); mt.cell(r_idx, 6).text = str(t_ds_h); mt.cell(r_idx, 7).text = str(t_ds_vd)
    mt.cell(r_idx, 8).text = str(t_tl_b); mt.cell(r_idx, 9).text = str(t_tl_h); mt.cell(r_idx, 10).text = str(t_tl_vd)
    
    # Bảng chữ ký
    doc.add_paragraph("\n")
    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_text(sig_table.cell(0, 0).paragraphs[0], "DUYỆT CỦA LÃNH ĐẠO", bold=True)
    
    sig_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_text(sig_table.cell(0, 1).paragraphs[0], "DUYỆT CỦA TỔ TRƯỞNG", bold=True)
    
    sig_table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_text(sig_table.cell(0, 2).paragraphs[0], "GIÁO VIÊN RA ĐỀ\n\n\n\n", bold=True)
    style_text(sig_table.cell(0, 2).paragraphs[0], f"{info['teacher']}", bold=True)
    
    # === TRANG 2: ĐỀ KIỂM TRA CHÍNH THỨC ===
    doc.add_page_break()
    header_exam = doc.add_table(rows=1, cols=2)
    header_exam.width = Inches(6)
    c_left = header_exam.cell(0, 0).paragraphs[0]
    style_text(c_left, f"{info['school'].upper()}\n", bold=True)
    style_text(c_left, f"GV: {info['teacher']}", italic=True)
    c_right = header_exam.cell(0, 1).paragraphs[0]
    c_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_text(c_right, f"ĐỀ KIỂM TRA MÔN {info['subject'].upper()}\n", bold=True)
    style_text(c_right, f"Thời gian làm bài: {info['time']} phút")

    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    list_tn = [q for q in exam_list if q['Loai'] == 'Trắc nghiệm']
    list_ds = [q for q in exam_list if q['Loai'] == 'Đúng/Sai']
    list_tl = [q for q in exam_list if q['Loai'] == 'Tự luận']

    if list_tn:
        doc.add_paragraph("\nPHẦN I. CÂU TRẮC NGHIỆM NHIỀU PHƯƠNG ÁN LỰA CHỌN.").runs[0].bold = True
        for i, q in enumerate(list_tn):
            p = doc.add_paragraph()
            style_text(p, f"Câu {i+1}: ", bold=True)
            style_text(p, str(q['Noi_dung']))
            for label in ['A', 'B', 'C', 'D']:
                opt_p = doc.add_paragraph()
                opt_p.paragraph_format.left_indent = Inches(0.3)
                style_text(opt_p, f"{label}. ", bold=True)
                style_text(opt_p, str(q['Options'][label]))

    if list_ds:
        doc.add_paragraph("\nPHẦN II. CÂU TRẮC NGHIỆM ĐÚNG SAI.").runs[0].bold = True
        for i, q in enumerate(list_ds):
            p = doc.add_paragraph()
            style_text(p, f"Câu {i+1}: ", bold=True)
            style_text(p, str(q['Noi_dung']))
            for label, sub in zip(['A', 'B', 'C', 'D'], ['a', 'b', 'c', 'd']):
                if q['Options'][label]:
                    opt_p = doc.add_paragraph()
                    opt_p.paragraph_format.left_indent = Inches(0.3)
                    style_text(opt_p, f"{sub}) ", bold=True)
                    style_text(opt_p, str(q['Options'][label]))

    if list_tl:
        doc.add_paragraph("\nPHẦN III. CÂU TRẮC NGHIỆM TRẢ LỜI NGẮN / TỰ LUẬN.").runs[0].bold = True
        for i, q in enumerate(list_tl):
            p = doc.add_paragraph()
            style_text(p, f"Câu {i+1}: ", bold=True)
            style_text(p, str(q['Noi_dung']))
            doc.add_paragraph("\n\n")

    # === TRANG 3: ĐÁP ÁN ===
    doc.add_page_break()
    doc.add_paragraph("BẢNG ĐÁP ÁN & HƯỚNG DẪN CHẤM").runs[0].bold = True
    
    if list_tn:
        doc.add_paragraph("Phần I. Trắc nghiệm").runs[0].italic = True
        ans_table = doc.add_table(rows=2, cols=len(list_tn))
        ans_table.style = 'Table Grid'
        for i, q in enumerate(list_tn):
            ans_table.cell(0, i).text = str(i+1)
            ans_table.cell(1, i).text = str(q['Correct'])

    if list_ds:
        doc.add_paragraph("\nPhần II. Đúng/Sai").runs[0].italic = True
        ans_table_ds = doc.add_table(rows=5, cols=len(list_ds) + 1)
        ans_table_ds.style = 'Table Grid'
        ans_table_ds.cell(0, 0).text = "Câu"
        ans_table_ds.cell(1, 0).text = "Ý a"; ans_table_ds.cell(2, 0).text = "Ý b"
        ans_table_ds.cell(3, 0).text = "Ý c"; ans_table_ds.cell(4, 0).text = "Ý d"
        for i, q in enumerate(list_ds):
            ans_table_ds.cell(0, i+1).text = str(i+1)
            ans_arr = str(q['Correct']).replace(' ', '').split(',')
            for j in range(4):
                if j < len(ans_arr): ans_table_ds.cell(j+1, i+1).text = ans_arr[j]

    if list_tl:
        doc.add_paragraph("\nPhần III. Tự luận/Trả lời ngắn").runs[0].italic = True
        for i, q in enumerate(list_tl):
            doc.add_paragraph(f"Câu {i+1}: {q['Correct']}")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- UI CHÍNH ---
st.title("🎓 Hệ Thống Tạo Đề Đa Định Dạng & Ma Trận (CV7991)")

with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/3074/3074058.png", width=100)
    st.header("⚙️ Thông tin chung")
    coquan = st.text_input("Cơ quan chủ quản", "UBND XÃ VĨNH HẬU")
    school = st.text_input("Tên trường", "TRƯỜNG THCS ĐA PHƯỚC")
    teacher = st.text_input("Giáo viên ra đề", "Tống Phước Khải")
    subject = st.text_input("Môn học", "KHTN 9 (HÓA HỌC)")
    duration = st.number_input("Thời gian (phút)", value=40)
    
    st.divider()
    st.markdown("🔑 **Khóa API AI (Tùy chọn)**")
    api_key = st.text_input("Nhập Gemini API Key", type="password")

tab1, tab2, tab3 = st.tabs(["📁 1. Dữ liệu", "✨ 2. Sinh Câu Hỏi AI (Có Công Thức)", "🖨️ 3. Lọc & Xuất Đề"])

with tab1:
    st.subheader("Tải lên Ngân hàng câu hỏi (Excel)")
    uploaded_file = st.file_uploader("Cột: Chuong, Muc_do, Loai_cau_hoi, Noi_dung, A, B, C, D, Dap_an_dung", type=["xlsx"])
    if st.button("Tải vào hệ thống", type="primary"):
        if uploaded_file:
            new_df = pd.read_excel(uploaded_file)
            if 'Loai_cau_hoi' not in new_df.columns: new_df['Loai_cau_hoi'] = 'Trắc nghiệm'
            new_df['Loai_cau_hoi'] = new_df['Loai_cau_hoi'].fillna('Trắc nghiệm')
            st.session_state.db_df = pd.concat([st.session_state.db_df, new_df], ignore_index=True).drop_duplicates()
            st.success("Đã bổ sung vào Kho!")
            
    if not st.session_state.db_df.empty:
        st.info(f"📚 Tổng: **{len(st.session_state.db_df)}** câu")
        st.dataframe(st.session_state.db_df[['Loai_cau_hoi', 'Muc_do', 'Noi_dung']].head(5))
        if st.button("🗑️ Xóa sạch bộ nhớ", type="secondary"):
            st.session_state.db_df = pd.DataFrame(columns=['Chuong', 'Muc_do', 'Loai_cau_hoi', 'Noi_dung', 'A', 'B', 'C', 'D', 'Dap_an_dung'])
            st.rerun()

with tab2:
    st.subheader("Trợ lý AI - Tạo mới câu hỏi đa định dạng")
    st.caption("✨ Đã nâng cấp thuật toán: Tự động sử dụng định dạng Unicode cho công thức Toán, Lý, Hóa (H₂O, x²) để xuất file Word không bị lỗi font.")
    chapter_ai = st.text_input("Nhập tên Chủ đề / Chương (VD: Alkene, Nguồn nhiên liệu...):")
    col_tn, col_ds, col_tl = st.columns(3)
    ai_tn = col_tn.number_input("Số câu Trắc nghiệm", min_value=0, value=2)
    ai_ds = col_ds.number_input("Số câu Đúng/Sai", min_value=0, value=1)
    ai_tl = col_tl.number_input("Số câu Tự luận", min_value=0, value=1)
        
    if st.button("✨ XUẤT XƯỞNG CÂU HỎI MỚI", type="primary"):
        if not api_key: st.error("❌ Nhập API Key ở Menu trái trước.")
        elif not chapter_ai: st.error("❌ Nhập tên Chủ đề cần tạo.")
        else:
            with st.spinner('🤖 AI đang soạn thảo câu hỏi kèm công thức chuẩn...'):
                gen_data = generate_questions_with_ai(api_key, subject, chapter_ai, ai_tn, ai_ds, ai_tl)
                if gen_data:
                    new_ai_df = pd.DataFrame(gen_data)
                    st.session_state.db_df = pd.concat([st.session_state.db_df, new_ai_df], ignore_index=True)
                    st.success("🎉 Đã tạo thành công đa định dạng!")
                    st.dataframe(new_ai_df)

with tab3:
    if st.session_state.db_df.empty:
        st.warning("⚠️ Kho dữ liệu trống.")
    else:
        st.subheader("Bốc đề & Xuất bản Kèm Ma Trận")
        c1, c2, c3 = st.columns(3)
        co_tn = len(st.session_state.db_df[st.session_state.db_df['Loai_cau_hoi'] == 'Trắc nghiệm'])
        co_ds = len(st.session_state.db_df[st.session_state.db_df['Loai_cau_hoi'] == 'Đúng/Sai'])
        co_tl = len(st.session_state.db_df[st.session_state.db_df['Loai_cau_hoi'] == 'Tự luận'])
        
        lay_tn = c1.number_input(f"Câu Trắc nghiệm (Kho: {co_tn})", min_value=0, max_value=co_tn, value=min(12, co_tn))
        lay_ds = c2.number_input(f"Câu Đúng/Sai (Kho: {co_ds})", min_value=0, max_value=co_ds, value=min(2, co_ds))
        lay_tl = c3.number_input(f"Câu Tự luận (Kho: {co_tl})", min_value=0, max_value=co_tl, value=min(1, co_tl))

        if st.button("🚀 XUẤT FILE (MA TRẬN + ĐỀ THI + ĐÁP ÁN)", type="primary", use_container_width=True):
            with st.spinner('Đang lập Ma Trận và xuất file Word...'):
                exam_list = []
                if lay_tn > 0:
                    for _, row in st.session_state.db_df[st.session_state.db_df['Loai_cau_hoi'] == 'Trắc nghiệm'].sample(n=lay_tn).iterrows(): exam_list.append(shuffle_question(row))
                if lay_ds > 0:
                    for _, row in st.session_state.db_df[st.session_state.db_df['Loai_cau_hoi'] == 'Đúng/Sai'].sample(n=lay_ds).iterrows(): exam_list.append(shuffle_question(row))
                if lay_tl > 0:
                    for _, row in st.session_state.db_df[st.session_state.db_df['Loai_cau_hoi'] == 'Tự luận'].sample(n=lay_tl).iterrows(): exam_list.append(shuffle_question(row))
                
                info = {'coquan': coquan, 'school': school, 'teacher': teacher, 'subject': subject, 'time': duration}
                docx_file = export_full_exam(exam_list, info)
                
                st.balloons()
                st.success("✅ Đã tạo thành công BỘ HỒ SƠ KIỂM TRA hoàn chỉnh!")
                st.download_button(
                    label="📥 TẢI XUỐNG BỘ ĐỀ (.docx)",
                    data=docx_file,
                    file_name=f"Ho_So_Kiem_Tra_{subject.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
